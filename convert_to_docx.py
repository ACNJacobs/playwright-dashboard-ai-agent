# Copyright (c) 2026 Ton Jacobs. All rights reserved.
# This file is part of the Playwright Test Suite.

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import re

import sys

# Lees het markdown bestand
input_file = sys.argv[1] if len(sys.argv) > 1 else 'PLAYWRIGHT_HANDLEIDING.md'
output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.md', '.docx')

with open(input_file, 'r', encoding='utf-8') as f:
    content = f.read()

# Maak een nieuw Word document
doc = Document()

# Pas de standaardstijl aan
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Definieer kopstijlen
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = 'Calibri'
    heading_font.bold = True
    if i == 1:
        heading_font.size = Pt(20)
        heading_font.color.rgb = RGBColor(0, 112, 192)
    elif i == 2:
        heading_font.size = Pt(16)
        heading_font.color.rgb = RGBColor(0, 112, 192)
    else:
        heading_font.size = Pt(13)
        heading_font.color.rgb = RGBColor(0, 112, 192)

# Code blok stijl toevoegen
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(9)
code_style.paragraph_format.space_after = Pt(6)
code_style.paragraph_format.space_before = Pt(6)

# Verwerk de markdown regel voor regel
lines = content.split('\n')
i = 0
in_code_block = False
code_lines = []
in_table = False
table_lines = []

while i < len(lines):
    line = lines[i]
    
    # Code blokken
    if line.startswith('```'):
        if in_code_block:
            # Einde code blok
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text, style='CodeBlock')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            # Achtergrondkleur voor code blok (grijs)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            p._p.get_or_add_pPr().append(shading_elm)
            in_code_block = False
            code_lines = []
        else:
            # Begin code blok
            in_code_block = True
        i += 1
        continue
    
    if in_code_block:
        code_lines.append(line)
        i += 1
        continue
    
    # Tabellen (markdown tabel syntax)
    if '|' in line and not line.startswith('#') and not line.startswith('```'):
        if not in_table:
            in_table = True
            table_lines = []
        table_lines.append(line)
        i += 1
        # Check of volgende regel ook een tabel is
        if i < len(lines) and '|' not in lines[i]:
            in_table = False
            # Verwerk tabel
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')]
                cells = [c for c in cells if c]  # Verwijder lege cellen
                if cells and not all(c.replace('-', '') == '' for c in cells):  # Skip scheidingsregel
                    rows.append(cells)
            
            if rows:
                num_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                for r_idx, row_cells in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_cells):
                        if c_idx < num_cols:
                            table.rows[r_idx].cells[c_idx].text = cell_text
                            # Eerste rij vet maken
                            if r_idx == 0:
                                for paragraph in table.rows[r_idx].cells[c_idx].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
            table_lines = []
        continue
    else:
        in_table = False
    
    # Horizontale lijn
    if line.strip() == '---':
        doc.add_paragraph()
        i += 1
        continue
    
    # Koppen
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
        i += 1
        continue
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
        i += 1
        continue
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
        i += 1
        continue
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
        i += 1
        continue
    
    # Lijstitems
    if line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        p = doc.add_paragraph(text, style='List Bullet')
        i += 1
        continue
    
    if re.match(r'^\d+\.\s', line.strip()):
        text = re.sub(r'^\d+\.\s', '', line.strip())
        p = doc.add_paragraph(text, style='List Number')
        i += 1
        continue
    
    # Vetgedrukte tekst binnen paragrafen
    if line.strip():
        p = doc.add_paragraph()
        # Verwerk **vet** en *cursief*
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line.strip())
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)
        i += 1
        continue
    
    # Lege regels
    if line.strip() == '':
        doc.add_paragraph()
    
    i += 1

# Sla het document op
doc.save(output_file)
print(f'DOCX bestand aangemaakt: {output_file}')
